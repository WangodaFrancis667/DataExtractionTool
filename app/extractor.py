import os
import pandas as pd
import pdfplumber
from docx import Document
import camelot
import magic  # python-magic to detect mime
from datetime import datetime
import re


def detect_type(path):
    m = magic.from_file(path, mime=True)
    if "pdf" in m:
        return "pdf"
    if "word" in m or path.lower().endswith(".docx"):
        return "docx"
    return None


def extract_from_docx(path, extract_text=True, extract_tables=True):
    doc = Document(path)
    txt = []
    tables = []
    if extract_text:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                txt.append(text)
    if extract_tables:
        for tbl in doc.tables:
            rows = []
            for r in tbl.rows:
                cells = [c.text.strip() for c in r.cells]
                rows.append(cells)
            df = pd.DataFrame(rows)
            tables.append(df)
    return "\n".join(txt), tables


def extract_from_pdf(path, extract_text=True, extract_tables=True):
    text = []
    tables = []
    # Try pdfplumber for text and tables first
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            if extract_text:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            if extract_tables:
                tbl = page.extract_table()
                if tbl:
                    df = pd.DataFrame(tbl[1:], columns=tbl[0])
                    tables.append(df)
    # If no tables found try Camelot (better for many table types)
    if extract_tables and not tables:
        try:
            camelot_tables = camelot.read_pdf(
                path, pages="all", flavor="lattice"
            )  # try lattice first
            if camelot_tables:
                for t in camelot_tables:
                    tables.append(t.df)
        except Exception:
            # fallback to stream mode
            try:
                camelot_tables = camelot.read_pdf(path, pages="all", flavor="stream")
                for t in camelot_tables:
                    tables.append(t.df)
            except Exception:
                pass
    return "\n".join(text), tables


def extract_fields_from_text(text, desired_fields):
    """
    Extract specific fields from text using pattern matching.
    Returns a dictionary with field names as keys and extracted values as values.
    """
    extracted_data = {}

    if not desired_fields or not text:
        return extracted_data

    for field in desired_fields:
        field = field.strip()
        if not field:
            continue

        # Common patterns for field extraction
        patterns = [
            rf"{re.escape(field)}\s*:?\s*([^\n\r,;]+)",  # Field: Value
            rf"{re.escape(field)}\s*[:-]\s*([^\n\r,;]+)",  # Field - Value or Field: Value
            rf"(?i){re.escape(field)}\s*:?\s*([^\n\r,;]+)",  # Case insensitive
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Take the first match and clean it up
                value = matches[0].strip()
                # Remove common trailing characters
                value = re.sub(r"[,;:]$", "", value).strip()
                extracted_data[field] = value
                break

        # If no match found, try to find the field name and get the text after it
        if field not in extracted_data:
            field_pattern = rf"(?i){re.escape(field)}"
            match = re.search(field_pattern, text)
            if match:
                # Get text after the field name (next 50 characters)
                start_pos = match.end()
                surrounding_text = text[start_pos : start_pos + 50].strip()
                # Extract meaningful value (stop at newline or common delimiters)
                value_match = re.match(r"[:\-\s]*([^\n\r,;]+)", surrounding_text)
                if value_match:
                    extracted_data[field] = value_match.group(1).strip()

    return extracted_data


def merge_tables_to_excel(tables, out_path, sheet_prefix="Table"):
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        for idx, df in enumerate(tables):
            name = f"{sheet_prefix}_{idx+1}"
            # sanitize sheet name length (Excel limit is 31 characters)
            sheet_name = name[:31]
            df.to_excel(writer, sheet_name=sheet_name, index=False)


def process_document(
    path,
    output_folder,
    extract_text=True,
    extract_tables=True,
    desired_fields=None,
    out_format="csv",
):
    doc_type = detect_type(path)
    if not doc_type:
        raise ValueError("Could not detect document type")

    if doc_type == "docx":
        text, tables = extract_from_docx(path, extract_text, extract_tables)
    else:
        text, tables = extract_from_pdf(path, extract_text, extract_tables)

    base_name = os.path.splitext(os.path.basename(path))[0]
    timestamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    out_basename = f"{base_name}_{timestamp}"

    os.makedirs(output_folder, exist_ok=True)

    # Extract specific fields if requested
    extracted_fields = {}
    if desired_fields and text:
        extracted_fields = extract_fields_from_text(text, desired_fields)

        # Save extracted fields to a separate file
        if extracted_fields:
            fields_path = os.path.join(output_folder, f"{out_basename}_fields.json")
            import json

            with open(fields_path, "w", encoding="utf-8") as f:
                json.dump(extracted_fields, f, indent=2, ensure_ascii=False)

            # Also create a CSV with the extracted fields
            fields_csv_path = os.path.join(output_folder, f"{out_basename}_fields.csv")
            fields_df = pd.DataFrame([extracted_fields])
            fields_df.to_csv(fields_csv_path, index=False)

    # Save full text if requested
    if extract_text:
        txt_path = os.path.join(output_folder, f"{out_basename}_text.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)

    # Process tables
    out_path = None
    if extract_tables and tables:
        if out_format == "excel":
            out_path = os.path.join(output_folder, f"{out_basename}_tables.xlsx")
            merge_tables_to_excel(tables, out_path)
        else:
            # CSV: create multiple CSV files for tables
            csv_paths = []
            for idx, df in enumerate(tables):
                p = os.path.join(output_folder, f"{out_basename}_table{idx+1}.csv")
                df.to_csv(p, index=False)
                csv_paths.append(p)
            out_path = csv_paths[0] if csv_paths else None

    # Return the most relevant output file path
    if extracted_fields:
        return os.path.join(output_folder, f"{out_basename}_fields.csv")
    elif out_path:
        return out_path
    elif extract_text:
        return txt_path
    else:
        return None


def list_outputs(output_folder):
    if not os.path.exists(output_folder):
        return []
    return sorted(os.listdir(output_folder), reverse=True)
